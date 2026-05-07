#!/usr/bin/env python3
"""
Generate a Word document (.docx) containing
setup and execution instructions for the
Rubrik CDM Pre-Upgrade Assessment Tool.

Covers Windows, Linux, and macOS.
Includes documentation for setup and run scripts.

Usage:
    python generate_docs.py

Output:
    Rubrik_CDM_Upgrade_Assessment_Guide.docx

Requires:
    pip install python-docx
"""

import sys
try:
    from docx import Document
    from docx.shared import (
        Inches, Pt, RGBColor
    )
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
except ImportError:
    print(
        "ERROR: python-docx is required.\n"
        "Install it with: pip install python-docx"
    )
    sys.exit(1)

from datetime import datetime


def add_code_block(doc, text):
    """Add a formatted code block paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.3)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(30, 30, 30)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "F2F2F2")
    shading.set(qn("w:val"), "clear")
    p.paragraph_format.element.get_or_add_pPr(
    ).append(shading)
    return p


def add_env_table(doc, rows):
    """Add environment variable table."""
    table = doc.add_table(
        rows=len(rows) + 1, cols=3,
    )
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    hdr[0].text = "Variable"
    hdr[1].text = "Description"
    hdr[2].text = "Example"
    for cell in hdr:
        for p in cell.paragraphs:
            p.runs[0].bold = True

    for idx, (var, desc, example) in enumerate(
        rows, 1
    ):
        table.rows[idx].cells[0].text = var
        table.rows[idx].cells[1].text = desc
        table.rows[idx].cells[2].text = example


def add_file_table(doc, rows):
    """Add a file inventory table."""
    table = doc.add_table(
        rows=len(rows) + 1, cols=3,
    )
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid Accent 1"

    hdr = table.rows[0].cells
    hdr[0].text = "File"
    hdr[1].text = "Platform"
    hdr[2].text = "Purpose"
    for cell in hdr:
        for p in cell.paragraphs:
            p.runs[0].bold = True

    for idx, (f, plat, purpose) in enumerate(
        rows, 1
    ):
        table.rows[idx].cells[0].text = f
        table.rows[idx].cells[1].text = plat
        table.rows[idx].cells[2].text = purpose


def generate_document():
    doc = Document()

     # ══════════════════════════════════════════
    #  COVER PAGE
    # ══════════════════════════════════════════

    # Add some spacing at the top
    for _ in range(3):
        doc.add_paragraph("")

    # Title
    title = doc.add_heading(
        "Rubrik CDM Pre-Upgrade Assessment Tool",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(
        "Setup & Execution Guide\n"
        "Windows | Linux | macOS"
    )
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # Version / Date
    doc.add_paragraph("")
    ver_para = doc.add_paragraph()
    ver_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = ver_para.add_run(
        "Generated: "
        + datetime.now().strftime(
            "%Y-%m-%d %H:%M:%S"
        )
    )
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(120, 120, 120)

    # Author attribution
    doc.add_paragraph("")
    author_para = doc.add_paragraph()
    author_para.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
    )
    run = author_para.add_run(
        "Built by Jacob Bryce — Advisory SE"
    )
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(50, 50, 50)

    # Spacer before disclaimer
    doc.add_paragraph("")
    doc.add_paragraph("")

    # ── Disclaimer Box ──
    # Using a single-cell table to create a
    # bordered box effect for the disclaimer
    disclaimer_table = doc.add_table(
        rows=1, cols=1
    )
    disclaimer_table.alignment = (
        WD_TABLE_ALIGNMENT.CENTER
    )
    cell = disclaimer_table.rows[0].cells[0]

    # Clear default paragraph and add formatted
    # disclaimer text
    cell.paragraphs[0].clear()

    # IMPORTANT header
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("IMPORTANT NOTICE")
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(180, 0, 0)

    # Disclaimer body
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p2.add_run(
        "THIS IS NOT A RUBRIK BUILT OR SUPPORTED "
        "TOOL. THIS CARRIES NO WARRANTIES OR "
        "SUPPORTABILITY BY RUBRIK OR ITS CREATOR."
    )
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(180, 0, 0)

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p3.add_run(
        "This tool was built using publicly "
        "available documentation for Rubrik RSC "
        "& CDM. It is intended to help facilitate "
        "the upgrade process for large and complex "
        "environments."
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(80, 80, 80)

    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p4.add_run(
        "PLEASE ALWAYS CHECK WITH THE LATEST "
        "RUBRIK DOCUMENTATION BEFORE PROCEEDING "
        "WITH ANY UPGRADE."
    )
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RGBColor(80, 80, 80)

    # Style the disclaimer cell with shading
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._element.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), "FFF3F3")
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)

    # Add cell borders
    tc_borders = OxmlElement("w:tcBorders")
    for border_name in [
        "w:top", "w:left",
        "w:bottom", "w:right"
    ]:
        border = OxmlElement(border_name)
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "8")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "CC0000")
        tc_borders.append(border)
    tc_pr.append(tc_borders)

    # Set cell width
    from docx.shared import Cm
    cell.width = Cm(16)

    doc.add_page_break()

    # ══════════════════════════════════════════
    #  TABLE OF CONTENTS
    # ══════════════════════════════════════════
    doc.add_heading("Table of Contents", level=1)
    toc_items = [
        "1. Overview",
        "2. Prerequisites",
        "3. Package Contents",
        "4. Quick Start — macOS / Linux",
        "5. Quick Start — Windows",
        "6. Manual Installation — macOS",
        "7. Manual Installation — Linux",
        "8. Manual Installation — Windows",
        "9. Configuration",
        "10. Running the Assessment",
        "11. Setup & Run Scripts Reference",
        "12. Understanding the Output",
        "13. Troubleshooting",
        "14. Environment Variable Reference",
    ]
    for item in toc_items:
        doc.add_paragraph(
            item, style="List Number"
        )
    doc.add_page_break()

    # ══════════════════════════════════════════
    #  1. OVERVIEW
    # ══════════════════════════════════════════
    doc.add_heading("1. Overview", level=1)
    doc.add_paragraph(
        "The Rubrik CDM Pre-Upgrade Assessment "
        "Tool evaluates Rubrik CDM clusters for "
        "upgrade readiness by connecting to "
        "Rubrik Security Cloud (RSC) via GraphQL "
        "API and optionally to CDM clusters "
        "directly via REST API."
    )
    doc.add_paragraph(
        "The tool performs the following checks:"
    )
    checks = [
        "CDM version End-of-Support (EOS) status "
        "and upgrade path validation",
        "Compatibility matrix validation (vSphere, "
        "ESXi, MSSQL, Oracle, Host OS)",
        "Workload inventory (VMs, databases, "
        "physical hosts, NAS)",
        "SLA compliance and replication topology",
        "Active live mount detection",
        "Node hardware and storage capacity",
        "Network configuration review",
        "Running jobs and queued tasks",
    ]
    for check in checks:
        doc.add_paragraph(
            check, style="List Bullet"
        )
    doc.add_paragraph(
        "Output is generated in HTML (interactive "
        "dashboard), JSON, and CSV formats."
    )
    doc.add_page_break()

    # ══════════════════════════════════════════
    #  2. PREREQUISITES
    # ══════════════════════════════════════════
    doc.add_heading("2. Prerequisites", level=1)

    doc.add_heading("All Platforms", level=2)
    prereqs = [
        "Python 3.8 or higher",
        "Network access to RSC "
        "(https://your-org.my.rubrik.com)",
        "RSC Service Account with read "
        "permissions (Client ID + Client Secret)",
        "Network access to CDM cluster node IPs "
        "(optional, for CDM direct API checks)",
        "CDM local admin credentials "
        "(optional, for CDM direct API)",
    ]
    for p in prereqs:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading(
        "RSC Service Account Setup", level=2
    )
    doc.add_paragraph(
        "1. Log in to RSC → Settings → "
        "Service Accounts"
    )
    doc.add_paragraph(
        "2. Create a new Service Account with "
        "the 'Read Only Admin' role"
    )
    doc.add_paragraph(
        "3. Download or copy the Client ID and "
        "Client Secret"
    )
    doc.add_paragraph(
        "4. Note your RSC URL "
        "(e.g., https://rubrik-gaia.my.rubrik.com)"
    )
    doc.add_page_break()

    # ══════════════════════════════════════════
    #  3. PACKAGE CONTENTS
    # ══════════════════════════════════════════
    doc.add_heading("3. Package Contents", level=1)
    doc.add_paragraph(
        "The distribution package contains "
        "the following files:"
    )

    add_file_table(doc, [
        (
            "setup.sh",
            "macOS / Linux",
            "Automated setup: installs Python, "
            "creates venv, installs dependencies, "
            "configures .env, validates setup",
        ),
        (
            "setup.bat",
            "Windows",
            "Automated setup: checks Python, "
            "creates venv, installs dependencies, "
            "configures .env, validates setup",
        ),
        (
            "run.sh",
            "macOS / Linux",
            "Run script: pre-flight checks, "
            "runs assessment, post-run summary, "
            "auto-opens HTML report",
        ),
        (
            "run.bat",
            "Windows",
            "Run script: pre-flight checks, "
            "runs assessment, post-run summary, "
            "auto-opens HTML report",
        ),
        (
            "main.py",
            "All",
            "Main assessment orchestrator",
        ),
        (
            "config.py",
            "All",
            "Configuration and environment "
            "variable management",
        ),
        (
            "rsc_client.py",
            "All",
            "RSC GraphQL + CDM REST API client",
        ),
        (
            "models.py",
            "All",
            "Data models for assessment results",
        ),
        (
            "cluster_discovery.py",
            "All",
            "Cluster discovery and enrichment",
        ),
        (
            "compatibility_matrix.py",
            "All",
            "Version compatibility validation",
        ),
        (
            "cdm_eos_data.json",
            "All",
            "CDM End-of-Support data",
        ),
        (
            "requirements.txt",
            "All",
            "Python package dependencies",
        ),
        (
            ".env.example",
            "All",
            "Configuration template",
        ),
        (
            "collectors/",
            "All",
            "Assessment collector modules",
        ),
        (
            "generate_docs.py",
            "All",
            "Generates this documentation",
        ),
    ])
    doc.add_page_break()

    # ══════════════════════════════════════════
    #  4. QUICK START — macOS / Linux
    # ══════════════════════════════════════════
    doc.add_heading(
        "4. Quick Start — macOS / Linux", level=1
    )
    doc.add_paragraph(
        "The fastest way to get started on macOS "
        "or Linux is to use the provided setup "
        "and run scripts. These handle all "
        "installation steps automatically."
    )

    doc.add_heading(
        "Step 1: Extract and Setup", level=2
    )
    add_code_block(
        doc,
        "# Extract the package\n"
        "unzip rubrik-cdm-upgrade-assessment.zip\n"
        "cd rubrik-cdm-upgrade-assessment\n\n"
        "# Run automated setup\n"
        "chmod +x setup.sh run.sh\n"
        "./setup.sh",
    )

    doc.add_paragraph(
        "The setup script will automatically:"
    )
    auto_steps = [
        "Detect your operating system "
        "(macOS or Linux)",
        "Check for Python 3.8+ and install "
        "if missing (via Homebrew on macOS or "
        "apt/dnf/yum on Linux)",
        "Create a Python virtual environment "
        "(.venv)",
        "Install all required Python packages",
        "Create .env from .env.example template",
        "Create output and logs directories",
        "Validate the entire setup",
    ]
    for s in auto_steps:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading(
        "Step 2: Configure Credentials", level=2
    )
    add_code_block(
        doc,
        "# Edit .env with your RSC credentials\n"
        "nano .env\n\n"
        "# Required values:\n"
        "# RSC_BASE_URL=https://your-org"
        ".my.rubrik.com\n"
        "# RSC_CLIENT_ID=client|your-id\n"
        "# RSC_CLIENT_SECRET=your-secret\n"
        "# TARGET_CDM_VERSION=9.1.0",
    )

    doc.add_heading(
        "Step 3: Run Assessment", level=2
    )
    add_code_block(doc, "./run.sh")

    doc.add_paragraph(
        "The run script will verify credentials, "
        "run the assessment, display a summary, "
        "and on macOS offer to open the HTML "
        "report in your default browser."
    )
    doc.add_page_break()

    # ══════════════════════════════════════════
    #  5. QUICK START — Windows
    # ══════════════════════════════════════════
    doc.add_heading(
        "5. Quick Start — Windows", level=1
    )
    doc.add_paragraph(
        "The fastest way to get started on "
        "Windows is to use the provided setup "
        "and run batch files."
    )

    doc.add_heading(
        "Step 1: Prerequisites", level=2
    )
    doc.add_paragraph(
        "Download and install Python 3.8+ from "
        "https://www.python.org/downloads/"
    )
    doc.add_paragraph(
        'IMPORTANT: Check "Add Python to PATH" '
        "during installation."
    )

    doc.add_heading(
        "Step 2: Extract and Setup", level=2
    )
    doc.add_paragraph(
        'Right-click the zip file → '
        '"Extract All..." then open Command '
        "Prompt or PowerShell in the extracted "
        "directory:"
    )
    add_code_block(
        doc,
        "cd rubrik-cdm-upgrade-assessment\n"
        "setup.bat",
    )

    doc.add_paragraph(
        "The setup script will automatically:"
    )
    win_steps = [
        "Verify Python is installed and in PATH",
        "Create a Python virtual environment",
        "Install all required Python packages",
        "Create .env from template",
        "Create output and logs directories",
        "Validate the entire setup",
    ]
    for s in win_steps:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading(
        "Step 3: Configure Credentials", level=2
    )
    add_code_block(
        doc,
        "notepad .env\n\n"
        "REM Required values:\n"
        "REM RSC_BASE_URL=https://your-org"
        ".my.rubrik.com\n"
        "REM RSC_CLIENT_ID=client|your-id\n"
        "REM RSC_CLIENT_SECRET=your-secret\n"
        "REM TARGET_CDM_VERSION=9.1.0",
    )

    doc.add_heading(
        "Step 4: Run Assessment", level=2
    )
    add_code_block(doc, "run.bat")

    doc.add_paragraph(
        "The run script will verify credentials, "
        "run the assessment, display a summary, "
        "and offer to open the HTML report in "
        "your default browser."
    )
    doc.add_page_break()
    # ══════════════════════════════════════════
    #  6. MANUAL INSTALLATION — macOS
    # ══════════════════════════════════════════
    doc.add_heading(
        "6. Manual Installation — macOS", level=1
    )
    doc.add_paragraph(
        "If you prefer to set up manually "
        "instead of using setup.sh:"
    )

    doc.add_heading(
        "Step 1: Install Python", level=2
    )
    doc.add_paragraph(
        "macOS typically includes Python 3. "
        "Verify:"
    )
    add_code_block(doc, "python3 --version")
    doc.add_paragraph(
        "If not installed, install via Homebrew:"
    )
    add_code_block(
        doc,
        "brew install python@3.11",
    )

    doc.add_heading(
        "Step 2: Extract and Install", level=2
    )
    add_code_block(
        doc,
        "unzip rubrik-cdm-upgrade-assessment.zip\n"
        "cd rubrik-cdm-upgrade-assessment\n"
        "python3 -m venv .venv\n"
        "source .venv/bin/activate\n"
        "pip install --upgrade pip\n"
        "pip install -r requirements.txt",
    )

    doc.add_heading(
        "Step 3: Configure and Run", level=2
    )
    add_code_block(
        doc,
        "cp .env.example .env\n"
        "nano .env   # Edit credentials\n"
        "mkdir -p output logs\n"
        "python main.py",
    )
    doc.add_page_break()

    # ══════════════════════════════════════════
    #  7. MANUAL INSTALLATION — Linux
    # ══════════════════════════════════════════
    doc.add_heading(
        "7. Manual Installation — Linux", level=1
    )

    doc.add_heading(
        "Ubuntu / Debian", level=2
    )
    add_code_block(
        doc,
        "sudo apt update\n"
        "sudo apt install -y python3 python3-pip "
        "python3-venv\n\n"
        "unzip rubrik-cdm-upgrade-assessment.zip\n"
        "cd rubrik-cdm-upgrade-assessment\n"
        "python3 -m venv .venv\n"
        "source .venv/bin/activate\n"
        "pip install --upgrade pip\n"
        "pip install -r requirements.txt\n\n"
        "cp .env.example .env\n"
        "nano .env\n"
        "mkdir -p output logs\n"
        "python main.py",
    )

    doc.add_heading(
        "RHEL / CentOS / Rocky", level=2
    )
    add_code_block(
        doc,
        "sudo dnf install -y python3 python3-pip\n\n"
        "unzip rubrik-cdm-upgrade-assessment.zip\n"
        "cd rubrik-cdm-upgrade-assessment\n"
        "python3 -m venv .venv\n"
        "source .venv/bin/activate\n"
        "pip install --upgrade pip\n"
        "pip install -r requirements.txt\n\n"
        "cp .env.example .env\n"
        "nano .env\n"
        "mkdir -p output logs\n"
        "python main.py",
    )
    doc.add_page_break()

    # ══════════════════════════════════════════
    #  8. MANUAL INSTALLATION — Windows
    # ══════════════════════════════════════════
    doc.add_heading(
        "8. Manual Installation — Windows",
        level=1,
    )

    doc.add_heading(
        "Step 1: Install Python", level=2
    )
    doc.add_paragraph(
        "Download Python 3.11+ from "
        "https://www.python.org/downloads/"
    )
    doc.add_paragraph(
        'IMPORTANT: Check "Add Python to PATH" '
        "during installation."
    )

    doc.add_heading(
        "Step 2: Extract and Install", level=2
    )
    add_code_block(
        doc,
        "Expand-Archive "
        "-Path rubrik-cdm-upgrade-assessment.zip "
        "-DestinationPath .\\rubrik-assessment\n"
        "cd rubrik-assessment\n\n"
        "python -m venv .venv\n"
        ".venv\\Scripts\\activate\n"
        "pip install --upgrade pip\n"
        "pip install -r requirements.txt",
    )

    doc.add_heading(
        "Step 3: Configure and Run", level=2
    )
    add_code_block(
        doc,
        "copy .env.example .env\n"
        "notepad .env\n"
        "mkdir output\n"
        "mkdir logs\n"
        "python main.py",
    )

    doc.add_heading(
        "Windows Troubleshooting", level=2
    )
    doc.add_paragraph(
        "If you get 'Execution Policy' errors:"
    )
    add_code_block(
        doc,
        "Set-ExecutionPolicy -Scope CurrentUser "
        "-ExecutionPolicy RemoteSigned",
    )
    doc.add_page_break()

    # ══════════════════════════════════════════
    #  9. CONFIGURATION
    # ══════════════════════════════════════════
    doc.add_heading("9. Configuration", level=1)
    doc.add_paragraph(
        "All configuration is done via the .env "
        "file. Copy .env.example to .env and "
        "edit with your values."
    )

    doc.add_heading(
        "Required Settings", level=2
    )
    add_env_table(doc, [
        (
            "RSC_BASE_URL",
            "RSC instance URL",
            "https://rubrik-gaia.my.rubrik.com",
        ),
        (
            "RSC_ACCESS_TOKEN_URI",
            "RSC token endpoint",
            "https://rubrik-gaia.my.rubrik.com"
            "/api/client_token",
        ),
        (
            "RSC_CLIENT_ID",
            "Service Account Client ID",
            "client|abc123...",
        ),
        (
            "RSC_CLIENT_SECRET",
            "Service Account Client Secret",
            "your-secret-here",
        ),
        (
            "TARGET_CDM_VERSION",
            "Target upgrade version",
            "9.1.0",
        ),
    ])

    doc.add_heading(
        "Optional Settings", level=2
    )
    add_env_table(doc, [
        (
            "INCLUDE_CLUSTERS",
            "Only assess these clusters",
            "cluster-1,cluster-2",
        ),
        (
            "EXCLUDE_CLUSTERS",
            "Skip these clusters",
            "lab-cluster",
        ),
        (
            "MAX_PARALLEL_CLUSTERS",
            "Parallel cluster assessments",
            "10",
        ),
        (
            "CDM_DIRECT_ENABLED",
            "Enable CDM direct API checks",
            "true",
        ),
        (
            "CDM_USERNAME",
            "CDM local admin username",
            "admin",
        ),
        (
            "CDM_PASSWORD",
            "CDM local admin password",
            "your-password",
        ),
        (
            "REPORT_FORMATS",
            "Output report formats",
            "html,json,csv",
        ),
        (
            "STREAMING_OUTPUT",
            "Disk-backed streaming for "
            "large environments",
            "false",
        ),
        (
            "COMPAT_DISPLAY_CAP",
            "Max servers to list by name "
            "in compatibility findings",
            "25",
        ),
    ])
    doc.add_page_break()
    # ══════════════════════════════════════════
    #  10. RUNNING THE ASSESSMENT
    # ══════════════════════════════════════════
    doc.add_heading(
        "10. Running the Assessment", level=1
    )

    doc.add_heading(
        "Using the Scripts (Recommended)", level=2
    )
    doc.add_paragraph(
        "The simplest way to run the assessment "
        "is with the provided run scripts:"
    )
    add_code_block(
        doc,
        "# macOS / Linux:\n"
        "./run.sh\n\n"
        "# Windows:\n"
        "run.bat",
    )
    doc.add_paragraph(
        "The run scripts perform pre-flight "
        "checks before starting:"
    )
    preflight = [
        "Verify virtual environment exists",
        "Verify .env file exists",
        "Check that credentials are not "
        "still placeholder values",
        "Activate the virtual environment",
        "Run the assessment (main.py)",
        "Display post-run summary with exit "
        "code interpretation",
        "Offer to open the HTML report in "
        "your browser",
    ]
    for p in preflight:
        doc.add_paragraph(p, style="List Bullet")

    doc.add_heading(
        "Running Directly", level=2
    )
    add_code_block(
        doc,
        "# macOS / Linux:\n"
        "source .venv/bin/activate\n"
        "python main.py\n\n"
        "# Windows:\n"
        ".venv\\Scripts\\activate\n"
        "python main.py",
    )

    doc.add_heading(
        "Assessment Phases", level=2
    )
    phases = [
        (
            "Phase 1: Discovery",
            "Discovers all CDM clusters "
            "registered in RSC via GraphQL API.",
        ),
        (
            "Phase 2: Enrichment",
            "Enriches cluster metadata including "
            "node IPs, connected state, and "
            "CDM direct API authentication.",
        ),
        (
            "Phase 3: Filtering",
            "Applies INCLUDE/EXCLUDE filters "
            "and skips disconnected clusters.",
        ),
        (
            "Phase 4: Assessment",
            "Runs all collectors against each "
            "cluster in parallel.",
        ),
        (
            "Phase 5: Reporting",
            "Generates HTML dashboard, JSON, "
            "and CSV reports.",
        ),
    ]
    for name, desc in phases:
        doc.add_heading(name, level=3)
        doc.add_paragraph(desc)

    doc.add_heading(
        "Exit Codes", level=2
    )
    exit_table = doc.add_table(rows=4, cols=2)
    exit_table.style = "Light Grid Accent 1"
    hdr = exit_table.rows[0].cells
    hdr[0].text = "Exit Code"
    hdr[1].text = "Meaning"
    for cell in hdr:
        for p in cell.paragraphs:
            p.runs[0].bold = True
    exits = [
        ("0", "No blockers found — clusters "
         "appear upgrade-ready"),
        ("1", "BLOCKERS found — upgrade should "
         "NOT proceed until resolved"),
        ("2", "Some clusters FAILED assessment "
         "— review errors before proceeding"),
    ]
    for idx, (code, meaning) in enumerate(
        exits, 1
    ):
        exit_table.rows[idx].cells[0].text = code
        exit_table.rows[idx].cells[1].text = meaning
    doc.add_page_break()

    # ══════════════════════════════════════════
    #  11. SETUP & RUN SCRIPTS REFERENCE
    # ══════════════════════════════════════════
    doc.add_heading(
        "11. Setup & Run Scripts Reference",
        level=1,
    )

    doc.add_heading("setup.sh (macOS / Linux)",
                    level=2)
    doc.add_paragraph(
        "Automated setup script for macOS and "
        "Linux. Handles the entire first-time "
        "setup process."
    )
    doc.add_paragraph("What it does:")
    setup_sh_steps = [
        "Detects OS (macOS vs Linux distro)",
        "Checks for Python 3.8+ — installs via "
        "Homebrew (macOS), apt (Debian/Ubuntu), "
        "or dnf/yum (RHEL/CentOS/Rocky) if missing",
        "Creates Python virtual environment (.venv)",
        "Installs all pip dependencies from "
        "requirements.txt",
        "Copies .env.example to .env if .env "
        "doesn't exist",
        "Detects if credentials are still "
        "placeholder values and prompts for edit",
        "Creates output/ and logs/ directories",
        "Sets executable permissions on run.sh",
        "Validates setup by checking imports, "
        "required files, and collectors directory",
        "Reports clear PASS/FAIL summary",
    ]
    for s in setup_sh_steps:
        doc.add_paragraph(s, style="List Bullet")
    add_code_block(
        doc,
        "chmod +x setup.sh\n"
        "./setup.sh",
    )

    doc.add_heading("setup.bat (Windows)",
                    level=2)
    doc.add_paragraph(
        "Automated setup script for Windows. "
        "Run from Command Prompt or PowerShell."
    )
    doc.add_paragraph("What it does:")
    setup_bat_steps = [
        "Verifies Python is installed and in PATH",
        "Creates Python virtual environment",
        "Installs all pip dependencies",
        "Copies .env.example to .env if needed",
        "Detects placeholder credentials",
        "Creates output and logs directories",
        "Validates setup with import checks",
        "Reports PASS/FAIL summary",
    ]
    for s in setup_bat_steps:
        doc.add_paragraph(s, style="List Bullet")
    add_code_block(doc, "setup.bat")

    doc.add_heading("run.sh (macOS / Linux)",
                    level=2)
    doc.add_paragraph(
        "Assessment execution script with "
        "pre-flight checks and post-run summary."
    )
    doc.add_paragraph("Pre-flight checks:")
    run_sh_checks = [
        "Verifies Python 3 is available",
        "Checks .venv virtual environment exists",
        "Verifies .env file exists",
        "Detects placeholder credentials",
        "Activates virtual environment",
    ]
    for s in run_sh_checks:
        doc.add_paragraph(s, style="List Bullet")
    doc.add_paragraph("Post-run features:")
    run_sh_post = [
        "Color-coded exit status banner "
        "(green=OK, red=blockers, "
        "yellow=failures)",
        "Displays path to latest output directory",
        "Shows path to HTML report and log file",
        "On macOS: offers to open HTML report "
        "in default browser",
    ]
    for s in run_sh_post:
        doc.add_paragraph(s, style="List Bullet")
    add_code_block(doc, "./run.sh")

    doc.add_heading("run.bat (Windows)",
                    level=2)
    doc.add_paragraph(
        "Windows assessment execution script."
    )
    doc.add_paragraph("Features:")
    run_bat_features = [
        "Checks virtual environment exists",
        "Verifies .env file and credentials",
        "Runs the assessment",
        "Color-coded status summary",
        "Locates latest output directory",
        "Offers to open HTML report in "
        "default browser",
    ]
    for s in run_bat_features:
        doc.add_paragraph(s, style="List Bullet")
    add_code_block(doc, "run.bat")
    doc.add_page_break()

    # ══════════════════════════════════════════
    #  12. UNDERSTANDING THE OUTPUT
    # ══════════════════════════════════════════
    doc.add_heading(
        "12. Understanding the Output", level=1
    )

    doc.add_heading("Output Directory", level=2)
    add_code_block(
        doc,
        "output/assessment_YYYYMMDD_HHMMSS/\n"
        "  assessment_report.html  "
        "# Interactive dashboard\n"
        "  assessment_report.json  "
        "# Machine-readable\n"
        "  all_issues.csv          "
        "# All findings\n"
        "  cluster_summary.csv     "
        "# Per-cluster summary\n"
        "  clusters/               "
        "# Per-cluster JSON\n"
        "    sh1-PaloAlto.json\n"
        "    sh2-Cork.json\n"
        "    ...",
    )

    doc.add_heading("HTML Dashboard", level=2)
    features = [
        "Executive summary cards (Discovered, "
        "Assessed, Blockers, Warnings)",
        "Cross-cluster issues table — all "
        "blockers and warnings in one view",
        "Per-cluster expandable cards with "
        "severity badges",
        "Blocker/warning findings auto-expand",
        "Informational findings collapsed "
        "to reduce noise",
        "Compatibility findings show individual "
        "server names and detected versions",
    ]
    for f in features:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("Severity Levels", level=2)
    sev_table = doc.add_table(rows=4, cols=3)
    sev_table.style = "Light Grid Accent 1"
    hdr = sev_table.rows[0].cells
    hdr[0].text = "Severity"
    hdr[1].text = "Meaning"
    hdr[2].text = "Action"
    for cell in hdr:
        for p in cell.paragraphs:
            p.runs[0].bold = True
    sevs = [
        ("BLOCKER", "Upgrade cannot proceed",
         "Must resolve before upgrade"),
        ("WARNING", "Potential issue",
         "Review and plan remediation"),
        ("INFO", "Informational",
         "No action required"),
    ]
    for idx, (s, m, a) in enumerate(sevs, 1):
        sev_table.rows[idx].cells[0].text = s
        sev_table.rows[idx].cells[1].text = m
        sev_table.rows[idx].cells[2].text = a
    doc.add_page_break()

    # ══════════════════════════════════════════
    #  13. TROUBLESHOOTING
    # ══════════════════════════════════════════
    doc.add_heading("13. Troubleshooting", level=1)

    issues = [
        (
            "RSC authentication fails",
            "Verify RSC_CLIENT_ID and "
            "RSC_CLIENT_SECRET in .env. Ensure "
            "the service account has 'Read Only "
            "Admin' role.",
        ),
        (
            "CDM direct API not available",
            "Ensure network connectivity to "
            "cluster node IPs (port 443). Set "
            "CDM_DIRECT_ENABLED=true. Verify "
            "CDM_USERNAME and CDM_PASSWORD.",
        ),
        (
            "Cloud clusters fail assessment",
            "Cloud-hosted CDM clusters (Azure, "
            "AWS) typically don't have direct "
            "network access. They will use "
            "RSC-only checks.",
        ),
        (
            "SSL/TLS warnings",
            "The urllib3 warning about LibreSSL "
            "on macOS is cosmetic and can be "
            "ignored.",
        ),
        (
            "Rate limiting or timeouts",
            "Increase API_TIMEOUT_SECONDS and "
            "API_MAX_RETRIES. Reduce "
            "MAX_PARALLEL_CLUSTERS.",
        ),
        (
            "setup.sh fails on Linux",
            "Ensure you have sudo privileges "
            "for package installation. Or "
            "install Python 3.8+ manually and "
            "re-run setup.sh.",
        ),
        (
            "setup.bat: 'python' not recognized",
            "Reinstall Python and check "
            "'Add Python to PATH'. Or use the "
            "full path to python.exe.",
        ),
        (
            "Windows Execution Policy error",
            "Run: Set-ExecutionPolicy -Scope "
            "CurrentUser -ExecutionPolicy "
            "RemoteSigned",
        ),
    ]
    for title, resolution in issues:
        doc.add_heading(title, level=3)
        doc.add_paragraph(resolution)

    doc.add_heading("Log Files", level=2)
    add_code_block(
        doc,
        "# macOS / Linux:\n"
        "cat logs/assessment_*.log | tail -100\n\n"
        "# Windows:\n"
        "Get-Content logs\\assessment_*.log "
        "-Tail 100",
    )
    doc.add_page_break()

    # ══════════════════════════════════════════
    #  14. ENVIRONMENT VARIABLE REFERENCE
    # ══════════════════════════════════════════
    doc.add_heading(
        "14. Environment Variable Reference",
        level=1,
    )

    all_vars = [
        ("RSC_BASE_URL", "RSC instance URL",
         "https://your-org.my.rubrik.com"),
        ("RSC_ACCESS_TOKEN_URI", "OAuth endpoint",
         "https://your-org.my.rubrik.com"
         "/api/client_token"),
        ("RSC_CLIENT_ID", "Client ID",
         "client|abc123"),
        ("RSC_CLIENT_SECRET", "Client secret",
         "(your secret)"),
        ("TARGET_CDM_VERSION", "Target version",
         "9.1.0"),
        ("INCLUDE_CLUSTERS", "Assess only these",
         "cluster-a,cluster-b"),
        ("EXCLUDE_CLUSTERS", "Skip these",
         "lab-cluster"),
        ("MAX_PARALLEL_CLUSTERS", "Parallelism",
         "10"),
        ("CDM_DIRECT_ENABLED", "CDM REST API",
         "true"),
        ("CDM_USERNAME", "CDM admin user",
         "admin"),
        ("CDM_PASSWORD", "CDM admin password",
         "(your password)"),
        ("REPORT_FORMATS", "Output formats",
         "html,json,csv"),
        ("OUTPUT_DIR", "Output directory",
         "./output"),
        ("LOG_DIR", "Log directory",
         "./logs"),
        ("STREAMING_OUTPUT", "Disk-backed mode",
         "false"),
        ("GRAPHQL_PAGE_SIZE", "Page size",
         "200"),
        ("API_TIMEOUT_SECONDS", "Timeout",
         "60"),
        ("API_MAX_RETRIES", "Retry attempts",
         "5"),
        ("COMPAT_DISPLAY_CAP",
         "Server name display cap",
         "25"),
    ]
    add_env_table(doc, all_vars)

    # ── Save ──
    filename = (
        "Rubrik_CDM_Upgrade_Assessment_Guide.docx"
    )
    doc.save(filename)
    print("Document generated: " + filename)
    return filename


if __name__ == "__main__":
    generate_document()        